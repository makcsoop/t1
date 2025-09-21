from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def create_site_report():
    """Создает отчет о состоянии сайта в формате Word"""
    
    # Создаем новый документ
    doc = Document()
    
    # Настраиваем стили документа
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    # Заголовок отчета
    title = doc.add_heading('Отчёт по состоянию сайта: https://example.com', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о периоде и компании
    doc.add_paragraph('Период: 01.01.2024 -- 31.12.2024')
    doc.add_paragraph('Управление: ООО "PingTower"')
    
    # Дата составления
    date_paragraph = doc.add_paragraph('Дата составления: ')
    date_paragraph.add_run('19.09.2025').bold = True
    
    doc.add_paragraph()  # Пустая строка
    
    # Создаем таблицу с метриками
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Наименования Метрики'
    hdr_cells[1].text = 'Описание'
    hdr_cells[2].text = 'Норма'
    hdr_cells[3].text = 'Результат'
    
    # Делаем заголовки жирными
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Данные для таблицы
    metrics_data = [
        {
            'name': 'Uptime',
            'desc': 'Процент времени, когда сайт был онлайн за период (сутки/месяц)',
            'norm': '>99.9%',
            'result': '99.96%'
        },
        {
            'name': 'Время отклика',
            'desc': 'Время, за которое сервер отвечает на HTTP-запрос',
            'norm': '<200 мс',
            'result': '180 мс'
        },
        {
            'name': 'Скорость загрузки',
            'desc': 'Полное время загрузки страницы до полной интерактивности',
            'norm': '<3 сек',
            'result': '2.8 сек'
        },
        {
            'name': 'Доступность API',
            'desc': 'Процент успешных ответов API (без ошибок 5xx)',
            'norm': '100%',
            'result': '99.7%'
        },
        {
            'name': 'SSL-сертификат',
            'desc': 'Срок действия SSL-сертификата — предупреждение при приближении к истечению',
            'norm': '>30 дн',
            'result': '45 дней'
        },
        {
            'name': 'DNS-записи',
            'desc': 'Проверка корректности A/CNAME/MX-записей домена',
            'norm': 'Корректны',
            'result': 'Корректны'
        },
        {
            'name': 'LCP',
            'desc': 'Время загрузки самого большого элемента на экране',
            'norm': '<2.5 сек',
            'result': '4.2 сек'
        },
        {
            'name': 'INP',
            'desc': 'Задержка между действием пользователя и визуальным откликом',
            'norm': '≤200 мс',
            'result': '280 мс'
        },
        {
            'name': 'SI',
            'desc': 'Средняя скорость визуальной загрузки страницы',
            'norm': '<3400 мс',
            'result': '5200 мс'
        },
        {
            'name': 'TBT',
            'desc': 'Общее время, когда страница не отвечает на действия',
            'norm': '<300 мс',
            'result': '420 мс'
        },
        {
            'name': 'CLS',
            'desc': 'Степень визуальной стабильности',
            'norm': '≤0.1',
            'result': '0.22'
        },
        {
            'name': 'FCP',
            'desc': 'Время появления первого осмысленного контента',
            'norm': '<1800 мс',
            'result': '1600 мс'
        }
    ]
    
    # Добавляем данные в таблицу
    for metric in metrics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric['name']
        row_cells[1].text = metric['desc']
        row_cells[2].text = metric['norm']
        
        # Результат делаем жирным
        result_paragraph = row_cells[3].paragraphs[0]
        result_run = result_paragraph.add_run(metric['result'])
        result_run.bold = True
    
    # Настраиваем ширину колонок
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(1.0)
    
    doc.add_paragraph()  # Пустая строка
    
    # Заключительный текст
    conclusion = doc.add_paragraph()
    conclusion.add_run('Сайт демонстрирует высокую доступность и стабильность на уровне. Однако ')
    conclusion.add_run('наблюдается ряд критических проблем на уровне пользовательского ')
    conclusion.add_run('интерфейса и фронтенд-производительности.\n\n')
    conclusion.add_run('Для повышения общего качества обслуживания пользователей и улучшения ')
    conclusion.add_run('позиций в поисковых системах рекомендуется в приоритетном порядке ')
    conclusion.add_run('провести работы по оптимизации фронтенда и стабилизации API.')
    
    # Сохраняем документ
    filename = f"Отчёт по состоянию сайта_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"Отчет сохранен в файл: {filename}")
    
    return filename

def create_parametrized_report(site_url, period, company, report_date, metrics_data=None):
    """Создает параметризованный отчет"""
    
    if metrics_data is None:
        metrics_data = [
            {
                'name': 'Uptime',
                'desc': 'Процент времени, когда сайт был онлайн за период (сутки/месяц)',
                'norm': '>99.9%',
                'result': '99.96%'
            },
            # ... остальные метрики по умолчанию
        ]
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    # Заголовок
    title = doc.add_heading(f'Отчёт по состоянию сайта: {site_url}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация
    doc.add_paragraph(f'Период: {period}')
    doc.add_paragraph(f'Управление: {company}')
    
    date_paragraph = doc.add_paragraph('Дата составления: ')
    date_paragraph.add_run(report_date).bold = True
    
    doc.add_paragraph()
    
    # Таблица
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    headers = ['Наименования Метрики', 'Описание', 'Норма', 'Результат']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for metric in metrics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric['name']
        row_cells[1].text = metric['desc']
        row_cells[2].text = metric['norm']
        
        result_paragraph = row_cells[3].paragraphs[0]
        result_run = result_paragraph.add_run(metric['result'])
        result_run.bold = True
    
    # Настройка ширины колонок
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(1.0)
    
    doc.add_paragraph()
    
    # Заключение
    conclusion = doc.add_paragraph()
    conclusion.add_run('Сайт демонстрирует высокую доступность и стабильность на уровне. Однако ')
    conclusion.add_run('наблюдается ряд критических проблем на уровне пользовательского ')
    conclusion.add_run('интерфейса и фронтенд-производительности.\n\n')
    conclusion.add_run('Для повышения общего качества обслуживания пользователей и улучшения ')
    conclusion.add_run('позиций в поисковых системах рекомендуется в приоритетном порядке ')
    conclusion.add_run('провести работы по оптимизации фронтенда и стабилизации API.')
    
    filename = f"Отчёт_{site_url.split('//')[-1]}_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save("reports/"+filename)
    print(f"Параметризованный отчет сохранен в файл: {filename}")
    
    return filename

if __name__ == "__main__":
    # Создаем стандартный отчет
    report_file = create_site_report()
    
    # Пример создания параметризованного отчета
    custom_metrics = [
        {
            'name': 'Uptime', 
            'desc': 'Время доступности сайта', 
            'norm': '>99.9%', 
            'result': '99.98%'
        },
        {
            'name': 'Время отклика', 
            'desc': 'Скорость ответа сервера', 
            'norm': '<200 мс', 
            'result': '150 мс'
        },
    ]
    
    custom_report = create_parametrized_report(
        site_url="https://my-site.com",
        period="01.06.2024 -- 31.12.2024",
        company='ООО "WebMonitor"',
        report_date="20.09.2025",
        metrics_data=custom_metrics
    )